from celery import shared_task
from django.conf import settings
from .models import FileUpload, ActivityLog
import os
import re
import logging

logger = logging.getLogger(__name__)

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. .docx files will not be processed.")

@shared_task
def process_file_word_count(file_upload_id):
    """
    Celery task to count words in uploaded file
    """
    try:
        file_upload = FileUpload.objects.get(id=file_upload_id)
        file_path = file_upload.file.path
        
        logger.info(f"Processing file: {file_path}")
        
        # Check if file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Count words based on file type
        if file_upload.file_type == '.txt':
            word_count = count_words_txt(file_path)
        elif file_upload.file_type == '.docx' and DOCX_AVAILABLE:
            word_count = count_words_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type or missing dependencies: {file_upload.file_type}")
        
        # Update file upload record
        file_upload.word_count = word_count
        file_upload.status = 'completed'
        file_upload.save()
        
        # Log activity
        ActivityLog.objects.create(
            user=file_upload.user,
            action='file_processed',
            metadata={
                'filename': file_upload.filename,
                'word_count': word_count,
                'file_size': file_upload.file_size,
                'processing_status': 'success'
            }
        )
        
        logger.info(f"File processed successfully. Word count: {word_count}")
        return {
            'status': 'success',
            'word_count': word_count,
            'file_id': file_upload_id
        }
        
    except FileUpload.DoesNotExist:
        logger.error(f"FileUpload with id {file_upload_id} not found")
        return {'status': 'error', 'message': 'File not found'}
        
    except Exception as e:
        logger.error(f"Error processing file {file_upload_id}: {str(e)}")
        
        # Update status to failed
        try:
            file_upload = FileUpload.objects.get(id=file_upload_id)
            file_upload.status = 'failed'
            file_upload.save()
            
            # Log failure
            ActivityLog.objects.create(
                user=file_upload.user,
                action='file_processing_failed',
                metadata={
                    'filename': file_upload.filename,
                    'error': str(e)
                }
            )
        except:
            pass
            
        return {'status': 'error', 'message': str(e)}


def count_words_txt(file_path):
    """Count words in a .txt file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                content = file.read()
        except:
            with open(file_path, 'r', encoding='cp1252') as file:
                content = file.read()
    
    words = re.findall(r'\b\w+\b', content)
    return len(words)


def count_words_docx(file_path):
    """Count words in a .docx file"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required to process .docx files")
    
    doc = Document(file_path)
    
    word_count = 0
    for paragraph in doc.paragraphs:
        words = re.findall(r'\b\w+\b', paragraph.text)
        word_count += len(words)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                words = re.findall(r'\b\w+\b', cell.text)
                word_count += len(words)
    
    return word_count